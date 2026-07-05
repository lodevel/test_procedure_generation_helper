"""Editor full-report export reuses the shared project_services engine.

Renders a real .docx from a selected test folder via workflow_editor.core.full_report
-> project_services.report_export, with NO bundle (extraction returns empty, the doc
still renders the metadata). Proves the cross-tree shared-export plumbing.
"""
import os
import sys
from pathlib import Path

import pytest

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
_REPO = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(_REPO / "src"))                                  # project_services
sys.path.insert(0, str(_REPO / "external" / "test_procedure_generation_helper"))  # workflow_editor

# The walk-up assumes the submodule layout; a STANDALONE editor checkout has no
# host src/ above it, so the shared report engine is truthfully absent -> skip.
pytest.importorskip("project_services")


def test_export_full_report_renders_docx(tmp_path):
    from docx import Document
    from workflow_editor.core import full_report
    from project_services import report_export

    # A template with Jinja placeholders the shared context fills.
    tpl = tmp_path / "tpl.docx"
    d = Document()
    d.add_paragraph("Report title: {{ title }}")
    d.add_paragraph("Test count: {{ tests|length }}")
    d.save(str(tpl))

    # A project with one test folder holding a procedure.json.
    proj = tmp_path / "proj"
    test_dir = proj / "tests" / "T1"
    test_dir.mkdir(parents=True)
    (test_dir / "procedure.json").write_text('{"meta": {"id": "T1"}}', encoding="utf-8")

    sidecar = report_export.create_default_sidecar()
    sidecar["title"] = "HELLO REPORT"
    out = tmp_path / "out.docx"

    result = full_report.export_full_report(
        proj, [test_dir], out, sidecar=sidecar, template_path=tpl)

    assert result == out and out.is_file()
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "HELLO REPORT" in text          # metadata from the sidecar
    assert "Test count: 1" in text          # one procedure made it into the context


def test_export_full_report_no_tests_raises(tmp_path):
    from workflow_editor.core import full_report
    empty = tmp_path / "empty"
    empty.mkdir()
    try:
        full_report.export_full_report(tmp_path, [empty], tmp_path / "o.docx",
                                       template_path=tmp_path / "missing.docx")
        assert False, "expected FullReportError"
    except full_report.FullReportError:
        pass
