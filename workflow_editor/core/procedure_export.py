"""Export the editor's canonical procedure (``procedure_text.md``) to a
shareable Markdown or Word document.

The workflow editor's source-of-truth artifact is ``procedure_text.md`` — the
canonical procedure DSL.  Its shape is::

    # <id / title>
    <one-line description>

    ## Meta
    key: value ...

    ## Equipment
    ID : type ...

    ## Steps
    1. <step>            (authored numbers, with gaps + @FOR/@IF directives)
    ...

    ## Expected
    {n} <op> value ...

Export is a *faithful rendering*, never a rewrite:

* ``#`` / ``##`` section headers become real document headings.
* The one-line description under the title becomes a normal paragraph.
* Every section body line (Meta / Equipment / Steps / Expected, including
  ``@FOR`` directives and the exact authored step numbers) is preserved
  **verbatim** in a monospace block.  Nothing is renumbered, reflowed or
  normalised, so identifiers, values and control-flow read exactly as
  authored — a Markdown ordered-list would silently renumber the steps and
  drop the authored gaps, which is why bodies are emitted as monospace.

The module is self-contained: it imports nothing from the parent app
(dependency direction is parent → editor) and no pack code.  Markdown export
is pure-stdlib; Word export lazily imports ``python-docx`` and raises
:class:`WordExportUnavailable` with install guidance when it is missing.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterator, Optional, Union

# Only levels 1–2 are structural in the procedure DSL (``# title`` /
# ``## section``). A deeper ``###``-prefixed body line is therefore preserved
# verbatim inside its section, never mis-split into a heading.
_HEADING_RE = re.compile(r"^(#{1,2})\s+(.*\S)\s*$")


class WordExportUnavailable(RuntimeError):
    """Raised when Word export is requested but ``python-docx`` is absent."""


# ---------------------------------------------------------------------------
#  Block model (single source feeding both renderers)
# ---------------------------------------------------------------------------

def _heading(line: str) -> Optional[tuple[int, str]]:
    """Return ``(level, title)`` for an ATX heading line, else ``None``."""
    m = _HEADING_RE.match(line)
    if not m:
        return None
    return len(m.group(1)), m.group(2)


# Block is one of:
#   ("heading", level:int, title:str)
#   ("prose",   lines:list[str])      # description paragraph under the title
#   ("body",    lines:list[str])      # verbatim section body (Meta/Steps/...)
Block = Union[tuple[str, int, str], tuple[str, list]]


def iter_blocks(text: str) -> Iterator[Block]:
    """Yield document blocks in order.

    The first non-heading run (the description under the ``# Title``, before
    any ``##`` section) is emitted as ``prose``; every later section body is
    emitted verbatim as ``body``.
    """
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i, n = 0, len(lines)
    seen_section = False
    while i < n:
        h = _heading(lines[i])
        if h:
            level, title = h
            if level >= 2:
                seen_section = True
            yield ("heading", level, title)
            i += 1
            continue
        run: list[str] = []
        while i < n and _heading(lines[i]) is None:
            run.append(lines[i])
            i += 1
        # Trim surrounding blank lines so blocks don't carry the blank
        # separators that sit between sections.
        while run and not run[0].strip():
            run.pop(0)
        while run and not run[-1].strip():
            run.pop()
        if not run:
            continue
        yield ("prose" if not seen_section else "body", run)


def first_title(text: str) -> Optional[str]:
    """Return the first ``# `` heading's text (the procedure id), if any."""
    for line in text.replace("\r\n", "\n").split("\n"):
        m = _HEADING_RE.match(line)
        if m and len(m.group(1)) == 1:
            return m.group(2)
    return None


# ---------------------------------------------------------------------------
#  Markdown
# ---------------------------------------------------------------------------

def _body_fence(lines: list) -> str:
    """Return a backtick fence longer than any backtick run in *lines*.

    Guards against a body line that is itself ```` ``` ```` (which would
    otherwise close the fenced block early and spill the rest of the section
    into raw Markdown).
    """
    longest = 0
    for ln in lines:
        run = 0
        for ch in ln:
            if ch == "`":
                run += 1
                longest = max(longest, run)
            else:
                run = 0
    return "`" * max(3, longest + 1)


def to_markdown(text: str) -> str:
    """Render the canonical procedure text to a self-contained Markdown doc.

    Headings stay headings; the description stays a paragraph; section bodies
    are wrapped in fenced ``text`` blocks so step numbers, ``@FOR`` directives
    and identifiers render byte-faithfully in any Markdown viewer.
    """
    parts: list[str] = []
    for block in iter_blocks(text):
        if block[0] == "heading":
            _, level, title = block
            parts.append("#" * level + " " + title)
            parts.append("")
        elif block[0] == "prose":
            parts.append("\n".join(block[1]))
            parts.append("")
        else:  # body — verbatim, fenced (fence sized so no line can close it)
            fence = _body_fence(block[1])
            parts.append(fence + "text")
            parts.extend(block[1])
            parts.append(fence)
            parts.append("")
    return "\n".join(parts).rstrip("\n") + "\n"


def write_markdown(text: str, out_path: Union[str, Path]) -> Path:
    """Render *text* and write it to *out_path* (UTF-8). Returns the path."""
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(to_markdown(text), encoding="utf-8")
    return out


# ---------------------------------------------------------------------------
#  Word (.docx)
# ---------------------------------------------------------------------------

def write_word(
    text: str,
    out_path: Union[str, Path],
    *,
    title: Optional[str] = None,
) -> Path:
    """Render *text* to a ``.docx`` at *out_path*. Returns the path.

    Lazily imports ``python-docx``; raises :class:`WordExportUnavailable`
    with install guidance if it is not installed (the editor still launches
    and Markdown export still works without it).
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ModuleNotFoundError as e:  # pragma: no cover - environment dependent
        raise WordExportUnavailable(
            "Word export needs the 'python-docx' package, which is not "
            "installed in this environment.\n\n"
            "Install it, then try again:\n"
            "    pip install python-docx\n\n"
            "(Markdown export works without it.)"
        ) from e

    doc = Document()
    doc.core_properties.title = title or first_title(text) or "Test Procedure"

    for block in iter_blocks(text):
        if block[0] == "heading":
            _, level, htitle = block
            doc.add_heading(htitle, level=min(level, 9))
        elif block[0] == "prose":
            doc.add_paragraph("\n".join(block[1]))
        else:  # body — one verbatim monospace paragraph per source line
            for line in block[1]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
